"""`.docx` 를 **문서 순서대로** 블록 목록으로 읽는다.

`python-docx` 의 `Document.paragraphs` 와 `Document.tables` 는 **서로 다른 목록**이라,
둘을 따로 읽으면 표가 어느 제목 밑에 있었는지가 사라진다. 작업표준서는 내용의
대부분이 표에 있으므로(작업순서표·관리기준표·이상조치표) 그렇게 읽으면 청크가
제목만 남고 알맹이가 빠진다.

그래서 XML 본문을 직접 순회해 `<w:p>` 와 `<w:tbl>` 을 **원래 순서대로** 낸다.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

#: 제목 스타일 → 깊이. 표지·머리말 스타일(`Cover Title` 등)은 제목이 아니다.
_HEADING = re.compile(r"^Heading (\d)$")

#: 셀 안의 줄바꿈. 표를 한 줄로 눕힐 때 쓰며, 원문에 줄이 나뉘어 있었다는 사실은 남긴다.
_CELL_BREAK = " / "


@dataclass
class Block:
    """문서의 한 덩어리. `kind` 는 `heading` · `text` · `table` 중 하나다."""

    kind: str
    text: str = ""
    level: int = 0
    rows: list[list[str]] = field(default_factory=list)

    def is_empty(self) -> bool:
        return not self.text.strip() and not self.rows


def _cell_text(cell) -> str:
    return _CELL_BREAK.join(
        line.strip() for line in cell.text.splitlines() if line.strip()
    )


def _table_rows(table: Table) -> list[list[str]]:
    """병합 셀 때문에 같은 값이 반복돼도 그대로 둔다 — 원문을 바꾸지 않는다."""
    return [[_cell_text(c) for c in row.cells] for row in table.rows]


def read_blocks(path: str) -> list[Block]:
    doc = Document(path)
    blocks: list[Block] = []

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para = Paragraph(child, doc)
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style is not None else ""
            m = _HEADING.match(style or "")
            if m:
                blocks.append(Block(kind="heading", text=text, level=int(m.group(1))))
            else:
                blocks.append(Block(kind="text", text=text))
        elif child.tag == qn("w:tbl"):
            rows = _table_rows(Table(child, doc))
            if any(any(c for c in r) for r in rows):
                blocks.append(Block(kind="table", rows=rows))

    return blocks


def render_table(rows: list[list[str]]) -> str:
    """표를 파이프 구분 텍스트로 만든다.

    마크다운 표로 만들지 않는다 — 헤더 구분선(`|---|`)은 검색에도 임베딩에도
    쓸모가 없고 토큰만 먹는다. 대신 **첫 행을 헤더로 보고 나머지 행 앞에 붙이지
    않는다**. 붙이면 같은 문구가 행 수만큼 반복돼 검색 점수가 왜곡된다.
    청크가 잘려 헤더가 떨어져 나갈 때만 `chunker` 가 헤더를 다시 붙인다.
    """
    return "\n".join(" | ".join(cell for cell in row) for row in rows)
