"""DOE UI 사용자 가이드 Word 문서 생성."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 여백 설정 ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── 스타일 헬퍼 ─────────────────────────────────────────────────
def set_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1):
    p = doc.add_heading('', level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {1: (31, 73, 125), 2: (70, 130, 180), 3: (47, 117, 181)}
    set_font(run, size=sizes.get(level, 12), bold=True, color=colors.get(level))
    return p

def add_para(doc, text, indent=0, bullet=False, color=None, bold=False, size=10.5):
    p = doc.add_paragraph()
    if bullet:
        p.style = doc.styles['List Bullet']
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def add_label_value(doc, label, value, indent=1.0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    r1 = p.add_run(f"{label}: ")
    set_font(r1, size=10, bold=True, color=(70, 130, 180))
    r2 = p.add_run(value)
    set_font(r2, size=10)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        # 헤더 배경색
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2E4057')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    # 데이터 행
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        fill = 'F0F4F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), fill)
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_box(doc, text, bg='E8F4FD', border='2E86AB'):
    """강조 박스."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=10, color=(30, 80, 120))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top','left','bottom','right']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '6')
        bdr.set(qn('w:space'), '4')
        bdr.set(qn('w:color'), border)
        pBdr.append(bdr)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), bg)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    return p


# ═══════════════════════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("실험계획법 (DOE) 시뮬레이션 시스템")
set_font(r, size=22, bold=True, color=(31, 73, 125))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("사용자 가이드")
set_font(r2, size=16, color=(70, 130, 180))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Formulation ML — DOE 메뉴 활용 안내")
set_font(r3, size=12, color=(128, 128, 128))

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("2026-06-18  |  Amorepacific R&D")
set_font(r4, size=10, color=(150, 150, 150))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 목차 안내
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "전체 메뉴 구성 및 활용 흐름", level=1)

add_box(doc,
    "DOE UI는 4단계 순서로 사용하는 것을 권장합니다:\n"
    "  1단계. 실험 설계  →  어떤 실험을 몇 번이나 할지 설계\n"
    "  2단계. ML 시뮬레이션  →  AI가 실험 결과를 즉시 예측\n"
    "  3단계. 최적 조건 탐색  →  최고 품질을 만드는 배합 자동 계산\n"
    "  4단계. 결과 분석  →  통계 분석으로 근거 문서화")

doc.add_paragraph()

# 흐름 테이블
add_table(doc,
    ["단계", "메뉴", "핵심 질문", "산출물"],
    [
        ["1단계", "실험 설계", "어떤 조합으로 몇 번 실험할까?", "설계 행렬(실험 계획서)"],
        ["2단계", "ML 시뮬레이션", "이 조건이면 품질이 어떻게 될까?", "예측 품질점수 테이블"],
        ["3단계", "최적 조건 탐색", "품질 최대화 배합비율은?", "최적 Sn/Ag/Cu% 조건"],
        ["4단계", "결과 분석", "어떤 인자가 가장 중요한가?", "ANOVA·파레토 통계 보고"],
    ],
    col_widths=[2.0, 3.5, 5.5, 4.5]
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. 실험 설계
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "1. 실험 설계", level=1)
add_para(doc, "실제 실험을 하기 전, AI 시뮬레이션에 투입할 '실험 계획서'를 작성하는 단계입니다. "
              "6가지 방법 중 목적에 맞는 방법을 선택합니다.", size=10.5)
doc.add_paragraph()

menus_design = [
    {
        "name": "완전요인설계 (Full Factorial)",
        "icon": "⚗️",
        "when": "모든 인자 조합을 빠짐없이 실험하고 싶을 때",
        "how": "인자(Sn%, Ag%, 용해온도 등) 활성화 → 수준 수 설정 → [설계 생성] 클릭",
        "use": "인자가 3~4개 이하이고 정확한 전수 분석이 필요할 때",
        "warn": "인자 5개 × 3수준 = 243번 실험 → 인자가 많으면 LHS 권장",
        "output": "모든 조합 설계 행렬",
    },
    {
        "name": "부분요인설계 (Fractional Factorial)",
        "icon": "✂️",
        "when": "실험 수를 줄이면서 핵심 효과는 놓치지 않을 때",
        "how": "인자 선택 → 해상도(Resolution) 설정 → [설계 생성]",
        "use": "인자 5개 이상, 예산·시간 제한이 있을 때",
        "warn": "교호작용 일부가 혼재(Aliasing)될 수 있음",
        "output": "완전요인의 1/2~1/4 축소 설계",
    },
    {
        "name": "CCD (Central Composite Design)",
        "icon": "🎯",
        "when": "품질이 포물선(2차 곡선) 반응을 보일 것 같을 때",
        "how": "인자 범위 설정 → [설계 생성]",
        "use": "Sn%가 너무 높거나 낮으면 품질이 저하되는 구간이 의심될 때",
        "warn": "축점(α)이 범위를 벗어날 수 있으므로 인자 범위 확인 필요",
        "output": "중심점·축점 포함 RSM 설계",
    },
    {
        "name": "Box-Behnken",
        "icon": "🛡️",
        "when": "모서리(극단값) 조건 실험이 위험하거나 불가능할 때",
        "how": "3인자 자동 구성 → [설계 생성]",
        "use": "고온+고압 동시 설정처럼 위험한 극단 조합을 피해야 할 때",
        "warn": "3인자만 활성화됨 (cu_pct, melt_time_min은 중심값 고정)",
        "output": "안전 구간 내 RSM 설계",
    },
    {
        "name": "다구치 (Taguchi)",
        "icon": "🔩",
        "when": "공정 노이즈에 강한(Robust) 최적 조건을 찾을 때",
        "how": "직교배열표 선택(L4~L18) → [설계 생성]",
        "use": "공급사 편차, 온도 변동 등 외부 노이즈가 있어도 품질이 안정된 조건 탐색",
        "warn": "S/N비 해석에 대한 별도 지식 필요",
        "output": "신호/잡음비 기반 직교 설계",
    },
    {
        "name": "LHS (Latin Hypercube Sampling)",
        "icon": "🌐",
        "when": "빠르게 넓은 탐색 공간을 골고루 커버하고 싶을 때",
        "how": "샘플 수 입력(기본 100) → [설계 생성]",
        "use": "초기 탐색, ML 학습 데이터 생성, 공간 전체를 고르게 샘플링",
        "warn": "각 구간에서 1번씩 샘플링되어 편향 없음 — 일반 랜덤보다 우수",
        "output": "균등 분포 100개 실험점",
    },
]

for m in menus_design:
    add_heading(doc, f"{m['icon']} {m['name']}", level=2)
    add_box(doc, f"언제 사용: {m['when']}")
    add_label_value(doc, "사용법", m['how'])
    add_label_value(doc, "활용 시나리오", m['use'])
    add_label_value(doc, "주의사항", m['warn'])
    add_label_value(doc, "생성 결과", m['output'])
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. ML 시뮬레이션
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "2. ML 시뮬레이션", level=1)
add_para(doc, "설계된 실험 조건을 실제로 진행하는 대신, 학습된 AI 모델이 품질을 즉시 예측합니다. "
              "실험 비용 없이 수백 가지 조건을 빠르게 검토할 수 있습니다.", size=10.5)
doc.add_paragraph()

menus_sim = [
    {
        "name": "배치 시뮬레이션",
        "icon": "⚡",
        "when": "설계한 100개 실험 조건의 품질점수를 한번에 예측할 때",
        "how": "샘플 수 입력(기본 100) → 모델 선택 → 공급사 선택 → [샘플 데이터 로드]",
        "result": "100행 예측 결과 테이블 + 품질 분포 히스토그램 + 상관관계 히트맵 + 3D 산점도 + 이상치 차트",
        "tip": "CSV 내보내기 버튼으로 결과를 Excel에서 추가 분석 가능",
    },
    {
        "name": "모델 선택",
        "icon": "🤖",
        "when": "4가지 AI 모델의 성능을 비교하고 최적 모델을 선택할 때",
        "how": "모델 목록 화면에서 RMSE / R² / MAPE 비교",
        "result": "모델별 성능 카드 (Ridge / RandomForest / GradientBoosting / XGBoost)",
        "tip": "정확도 우선: GradientBoosting(R²=0.627) / 해석 우선: Ridge",
    },
    {
        "name": "민감도 분석",
        "icon": "📡",
        "when": "어떤 인자 하나를 건드리면 품질이 얼마나 변하는지 확인할 때",
        "how": "[민감도 분석] 클릭",
        "result": "각 인자 ±10% 변화 시 품질 변화폭 막대 차트",
        "tip": "민감도 높은 인자 = 관리 우선순위, 낮은 인자 = 허용 범위 넓게 설정 가능",
    },
    {
        "name": "몬테카를로 시뮬레이션",
        "icon": "🎲",
        "when": "공급사 편차와 공정 변동이 있을 때 실제 불량률을 추정할 때",
        "how": "노이즈 수준 설정 → [시뮬레이션]",
        "result": "품질 분포, 불량률(%), 6시그마 Cpk 추정",
        "tip": "실제 공정 변동성을 반영한 현실적 품질 예측에 활용",
    },
]

for m in menus_sim:
    add_heading(doc, f"{m['icon']} {m['name']}", level=2)
    add_box(doc, f"언제 사용: {m['when']}")
    add_label_value(doc, "사용법", m['how'])
    add_label_value(doc, "결과 화면", m['result'])
    add_label_value(doc, "활용 팁", m['tip'])
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. 최적 조건 탐색
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "3. 최적 조건 탐색", level=1)
add_para(doc, "AI 최적화 알고리즘(scipy SLSQP + LHS)이 품질을 최대화하는 배합비율을 자동으로 찾아줍니다. "
              "인자 범위만 지정하면 Sn+Ag+Cu+Pb=100% 제약 조건 하에서 최적 조건을 계산합니다.", size=10.5)
doc.add_paragraph()

menus_opt = [
    {
        "name": "단일목적 최적화",
        "icon": "🏆",
        "when": "품질점수를 최대화하는 배합비율 한 가지를 도출할 때",
        "how": "인자 범위 설정 → 모델/공급사 선택 → [최적화 실행]",
        "result": "최적 Sn%, Ag%, Cu%, 예측 품질점수, 상위 5개 후보 조건",
        "tip": "실제 결과 예시: Sn=62.6%, Ag=3.4%, Cu=0.59% → 품질 93.5점",
    },
    {
        "name": "다목적 최적화",
        "icon": "⚖️",
        "when": "품질↑ + 원가↓ + 불량률↓ 여러 목표를 동시에 만족할 때",
        "how": "목표별 가중치 조정 → [최적화]",
        "result": "파레토 최적 해집합 시각화",
        "tip": "품질과 비용 간 트레이드오프 구간을 시각적으로 파악",
    },
    {
        "name": "반응표면 (RSM)",
        "icon": "🗺️",
        "when": "두 인자 조합에 따른 품질 지형도를 3D로 보고 싶을 때",
        "how": "X축·Y축 인자 선택 → [반응표면 생성]",
        "result": "3D 등고선 + 표면 플롯",
        "tip": "'봉우리(최대점)'가 어디 있는지 직관적으로 확인",
    },
    {
        "name": "제약 최적화",
        "icon": "🔒",
        "when": "현실적 제약(합계=100%, 특정 범위 내)을 반드시 지켜야 할 때",
        "how": "제약 조건 입력 → [최적화 실행]  (내부: /doe/optimize API 호출)",
        "result": "제약 충족 최적 배합 + 최적화 수렴 이력",
        "tip": "SLSQP 수렴 실패 시 LHS best 후보로 자동 전환",
    },
]

for m in menus_opt:
    add_heading(doc, f"{m['icon']} {m['name']}", level=2)
    add_box(doc, f"언제 사용: {m['when']}")
    add_label_value(doc, "사용법", m['how'])
    add_label_value(doc, "결과 화면", m['result'])
    add_label_value(doc, "활용 팁", m['tip'])
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. 결과 분석
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "4. 결과 분석", level=1)
add_para(doc, "시뮬레이션 결과를 통계 분석하여 '어떤 인자가 왜 중요한가'를 수치로 입증합니다. "
              "보고서·근거 자료 작성에 활용합니다.", size=10.5)
doc.add_paragraph()

menus_analysis = [
    {
        "name": "주효과도",
        "icon": "📈",
        "when": "각 인자를 개별로 변화시켰을 때 품질에 미치는 영향을 비교할 때",
        "how": "[샘플 로드] → 자동 표시",
        "result": "인자별 수준 vs 평균 품질점수 꺾은선 그래프",
        "interpret": "기울기가 클수록 해당 인자가 품질에 큰 영향 → 관리 우선순위",
    },
    {
        "name": "교호작용도",
        "icon": "🔀",
        "when": "두 인자가 함께 변할 때 나타나는 상호작용을 파악할 때",
        "how": "[샘플 로드] → 자동 표시",
        "result": "SN% × AG% 조합별 품질 변화 교차선 그래프",
        "interpret": "선이 평행 = 교호작용 없음 / 선이 교차 = 상호작용 있음 (복합 관리 필요)",
    },
    {
        "name": "ANOVA 분석",
        "icon": "📊",
        "when": "각 인자 효과의 통계적 유의성을 수치로 검증해야 할 때",
        "how": "[ANOVA 분석] 버튼 클릭",
        "result": "SS(제곱합) / df / MS / F값 / p-value / 기여율(%) 테이블",
        "interpret": "p<0.05(*): 통계적 유의 / p<0.01(**): 강한 효과 / p<0.001(***): 매우 강함\nns: 통계적으로 무의미 → 해당 인자 범위 완화 가능",
    },
    {
        "name": "파레토 차트",
        "icon": "📉",
        "when": "중요한 인자 20%를 선별해 집중 관리할 때",
        "how": "[파레토 분석] 버튼 클릭",
        "result": "효과 크기 내림차순 막대 + 누적 80% 기준선 오버레이",
        "interpret": "80% 기준선 왼쪽 인자 = 핵심 관리 대상\n통상 1~3개 인자가 품질 변동의 80% 설명",
    },
    {
        "name": "이력 비교 ★",
        "icon": "🔍",
        "when": "실제 생산 이력 데이터로 시뮬레이션 파라미터를 재보정할 때",
        "how": "이력 수 입력(기본 50) → [이력 비교 실행]",
        "result": "공급사별 실측 평균 품질·표준편차 테이블\n+ SUPPLIER_EFFECTS 재보정 제안값\n+ 공급사별 품질 막대차트 + 이력 산점도",
        "interpret": "현재 설정값 vs 재보정값이 0.1 이상 차이나면 노란색 강조\n→ sample_generator.py의 SUPPLIER_EFFECTS 업데이트 권장",
    },
]

for m in menus_analysis:
    add_heading(doc, f"{m['icon']} {m['name']}", level=2)
    add_box(doc, f"언제 사용: {m['when']}")
    add_label_value(doc, "사용법", m['how'])
    add_label_value(doc, "결과 화면", m['result'])
    add_label_value(doc, "해석 방법", m['interpret'])
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. 권장 워크플로우
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "5. 권장 워크플로우", level=1)
add_para(doc, "처음 DOE 분석을 시작하는 경우 아래 6단계 순서를 따르는 것을 권장합니다.", size=10.5)
doc.add_paragraph()

steps = [
    ("STEP 1", "LHS로 100개 설계 생성",
     "실험 설계 > LHS 선택 → 샘플 수 100 → [설계 생성]\n목적: 탐색 공간 전체를 균등하게 커버하는 초기 후보 확보"),
    ("STEP 2", "배치 시뮬레이션 → 품질 분포 확인",
     "ML 시뮬레이션 > 배치 시뮬레이션 → [샘플 데이터 로드]\n목적: 현재 설계 범위에서 품질 분포·최대·최소 파악"),
    ("STEP 3", "ANOVA + 파레토로 핵심 인자 선별",
     "결과 분석 > ANOVA 분석 → p<0.05 인자 확인\n결과 분석 > 파레토 차트 → 80% 기준선 왼쪽 인자 집중 선택\n목적: 관리할 인자를 3개 이하로 좁히기"),
    ("STEP 4", "핵심 인자만 CCD로 정밀 설계",
     "실험 설계 > CCD → 핵심 인자 3개만 활성화 → [설계 생성]\n목적: RSM 분석을 위한 정밀 실험 계획서 작성"),
    ("STEP 5", "단일목적 최적화 → 최적 배합 도출",
     "최적 조건 탐색 > 단일목적 최적화 → [최적화 실행]\n목적: Sn+Ag+Cu+Pb=100% 제약 하 품질 최대화 조건 자동 계산"),
    ("STEP 6", "이력 비교 → 실데이터 검증 및 재보정",
     "결과 분석 > 이력 비교 → [이력 비교 실행]\n목적: 실생산 이력과 비교해 시뮬레이션 모델의 신뢰도 확인 및 공급사 효과 재보정"),
]

for step, title, detail in steps:
    p = doc.add_paragraph()
    r1 = p.add_run(f"[{step}]  ")
    set_font(r1, size=11, bold=True, color=(255, 255, 255))
    # step 배경 흉내 (bold blue)
    r1.font.color.rgb = RGBColor(31, 73, 125)
    r2 = p.add_run(title)
    set_font(r2, size=11, bold=True, color=(31, 73, 125))
    add_para(doc, detail, indent=0.5, size=10)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 6. 빠른 참조표
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "6. 빠른 참조표 — 목적별 메뉴 선택 가이드", level=1)
doc.add_paragraph()

add_table(doc,
    ["목적", "권장 메뉴", "소요 시간"],
    [
        ["처음 탐색 / 빠른 확인",    "실험 설계 > LHS + 배치 시뮬레이션",            "1~2분"],
        ["인자 중요도 파악",          "결과 분석 > ANOVA + 파레토",                   "2분"],
        ["최적 배합 도출",            "최적 조건 탐색 > 단일목적 최적화",              "30초"],
        ["실험 수 최소화",            "실험 설계 > 부분요인 or 다구치",               "3~5분"],
        ["2차 반응 포함 정밀 분석",   "실험 설계 > CCD + 반응표면",                  "5~10분"],
        ["노이즈 강건 조건 탐색",     "실험 설계 > 다구치 + 몬테카를로",             "5분"],
        ["실데이터 기반 검증/보정",   "결과 분석 > 이력 비교",                        "1분"],
        ["현실 제약 하 최적화",       "최적 조건 탐색 > 제약 최적화",                 "1분"],
    ],
    col_widths=[5.0, 7.0, 3.0]
)

doc.add_paragraph()
add_box(doc,
    "접속 주소:  http://localhost:8000/static/doe.html\n"
    "서버 실행:  python -m uvicorn app:app --port 8000\n"
    "API 문서:   http://localhost:8000/docs",
    bg='FFF9E6', border='F59E0B')

# ── 저장 ──────────────────────────────────────────────────────
output_path = r"C:\Users\amore\Desktop\26년 AP Claude Play Ground\04 Formulation ML\DOE_UI_사용자가이드.docx"
doc.save(output_path)
print(f"저장 완료: {output_path}")
