"""
Markdown → Word (공식 스마트공장 산출물 양식) 변환 스크립트
사용법: python scripts/md_to_docx.py
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 경로 ────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent.parent
INPUT_DIR = BASE_DIR / "docs" / "산출물"
OUTPUT_DIR = BASE_DIR / "docs" / "산출물" / "word"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ── 색상 ────────────────────────────────────────────────────────
COLOR_PRIMARY   = RGBColor(0x00, 0x47, 0x9D)   # 스마트공장 딥블루
COLOR_SECONDARY = RGBColor(0x00, 0x78, 0xD4)   # 밝은 블루
COLOR_ACCENT    = RGBColor(0x10, 0x74, 0x4E)   # 그린
COLOR_GRAY      = RGBColor(0x60, 0x60, 0x60)   # 회색
COLOR_LIGHT_BG  = RGBColor(0xF0, 0xF4, 0xFA)   # 연한 배경
COLOR_TABLE_HDR = RGBColor(0x00, 0x47, 0x9D)   # 테이블 헤더
COLOR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, rgb: RGBColor):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    hex_color = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, border_color="004799"):
    """셀 테두리 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_horizontal_rule(doc):
    """수평선 추가"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '00479D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_doc_styles(doc):
    """문서 기본 스타일 설정"""
    # 기본 폰트 (맑은 고딕)
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 여백
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)


def add_cover_page(doc, doc_id: str, title: str, meta: dict):
    """표지 페이지 생성"""
    # 상단 로고 영역
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = logo_para.add_run("스마트공장 보급·확산사업")
    run.font.name = '맑은 고딕'
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_GRAY
    run.font.bold = False

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 문서번호 배지
    id_para = doc.add_paragraph()
    id_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = id_para.add_run(f"[ {doc_id} ]")
    run.font.name = '맑은 고딕'
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_SECONDARY
    run.font.bold = True

    doc.add_paragraph()

    # 제목 (대형)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_clean = re.sub(r'\*\*|__', '', title).strip()
    run = title_para.add_run(title_clean)
    run.font.name = '맑은 고딕'
    run.font.size = Pt(26)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.bold = True

    doc.add_paragraph()

    # 구분선
    add_horizontal_rule(doc)
    doc.add_paragraph()

    # 프로젝트명
    proj_para = doc.add_paragraph()
    proj_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = proj_para.add_run("성분분석 데이터 기반 배합비율 최적화 ML 시스템")
    run.font.name = '맑은 고딕'
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_SECONDARY
    run.font.bold = False

    proj_para2 = doc.add_paragraph()
    proj_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = proj_para2.add_run("Formulation ML — Smart Factory AI/MES")
    run2.font.name = 'Calibri'
    run2.font.size = Pt(11)
    run2.font.color.rgb = COLOR_GRAY
    run2.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 문서 정보 테이블
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    info_rows = [
        ("문서번호", doc_id, "버전", meta.get("version", "V1.0")),
        ("작성일", meta.get("date", "2026-06-19"), "보안등급", "일반"),
        ("작성자", meta.get("author", "개발팀"), "승인자", "PM"),
        ("수행기관", "아모레퍼시픽", "사업명", "선도형 제조혁신 지원사업"),
        ("상태", "작성완료", "페이지", "—"),
    ]

    for r_idx, (k1, v1, k2, v2) in enumerate(info_rows):
        row = table.rows[r_idx]
        for c_idx, (text, is_key) in enumerate([(k1, True), (v1, False), (k2, True), (v2, False)]):
            cell = row.cells[c_idx]
            cell.text = text
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0] if para.runs else para.add_run(text)
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9)
            if is_key:
                set_cell_bg(cell, COLOR_PRIMARY)
                run.font.color.rgb = COLOR_WHITE
                run.font.bold = True
            else:
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            set_cell_border(cell)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 하단 기관명
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("아모레퍼시픽 제조사업부")
    run.font.name = '맑은 고딕'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY

    year_para = doc.add_paragraph()
    year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year_para.add_run("2026")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_GRAY

    # 페이지 나눔
    doc.add_page_break()


def add_header_footer(doc, doc_id: str, title: str):
    """헤더/푸터 추가"""
    section = doc.sections[0]

    # 헤더
    header = section.header
    header_para = header.paragraphs[0]
    header_para.clear()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run_left = header_para.add_run(f"{doc_id}  {title}")
    run_left.font.name = '맑은 고딕'
    run_left.font.size = Pt(8)
    run_left.font.color.rgb = COLOR_GRAY

    # 푸터 — 페이지 번호
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_f = footer_para.add_run("아모레퍼시픽 제조사업부  |  스마트공장 Formulation ML  |  ")
    run_f.font.name = '맑은 고딕'
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = COLOR_GRAY

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_page = footer_para.add_run()
    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = COLOR_GRAY


def parse_meta_from_md(lines: list[str]) -> dict:
    """Markdown 첫 몇 줄에서 문서 메타 파싱"""
    meta = {}
    for line in lines[:10]:
        if '버전' in line or 'V1.0' in line:
            m = re.search(r'V(\d+\.\d+)', line)
            if m:
                meta['version'] = f"V{m.group(1)}"
        if '작성일' in line:
            m = re.search(r'(\d{4}-\d{2}-\d{2})', line)
            if m:
                meta['date'] = m.group(1)
        if '작성자' in line:
            m = re.search(r'작성자[:\s*|]+(.+?)[\s*|]', line)
            if m:
                meta['author'] = m.group(1).strip()
    return meta


def add_styled_heading(doc, text: str, level: int):
    """스타일이 적용된 제목 추가"""
    text = re.sub(r'\*\*|__', '', text).strip()
    if not text:
        return

    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._p.get_or_add_pPr()
        # 위 공백
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(text)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        # 하단 테두리
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '00479D')
        pBdr.append(bottom)
        pPr.append(pBdr)

    elif level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY

    elif level == 3:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run("▶ " + text)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_ACCENT

    elif level >= 4:
        p = doc.add_paragraph()
        run = p.add_run("◆ " + text)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = COLOR_GRAY


def add_styled_paragraph(doc, text: str):
    """본문 단락 추가 (인라인 bold/italic 처리)"""
    if not text.strip():
        doc.add_paragraph()
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0)

    # bold(**text**), italic(*text*)
    pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            run = p.add_run(text[last:m.start()])
            run.font.name = '맑은 고딕'
            run.font.size = Pt(10)
        if m.group(1):  # **bold**
            run = p.add_run(m.group(1))
            run.font.bold = True
            run.font.name = '맑은 고딕'
            run.font.size = Pt(10)
        elif m.group(2):  # *italic*
            run = p.add_run(m.group(2))
            run.font.italic = True
            run.font.name = '맑은 고딕'
            run.font.size = Pt(10)
        elif m.group(3):  # `code`
            run = p.add_run(m.group(3))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        last = m.end()
    if last < len(text):
        run = p.add_run(text[last:])
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10)


def add_bullet(doc, text: str, level: int = 0):
    """불릿 항목 추가"""
    text = re.sub(r'^[-*+]\s+', '', text).strip()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    bullet = "•" if level == 0 else "◦"
    run = p.add_run(f"{bullet}  {text}")
    run.font.name = '맑은 고딕'
    run.font.size = Pt(10)


def add_numbered_item(doc, text: str, number: str):
    """번호 목록 항목"""
    text = re.sub(r'^\d+\.\s+', '', text).strip()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{number}  {text}")
    run.font.name = '맑은 고딕'
    run.font.size = Pt(10)


def add_code_block(doc, lines: list[str], lang: str = ""):
    """코드 블록 추가"""
    # 코드 블록 제목
    if lang:
        label_p = doc.add_paragraph()
        run = label_p.add_run(f"  {lang.upper()}")
        run.font.name = '맑은 고딕'
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = COLOR_WHITE
        label_p.paragraph_format.left_indent = Cm(0)
        # 배경색은 Word XML로
        pPr = label_p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '004799')
        pPr.append(shd)

    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        # 배경
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F4FA')
        pPr.append(shd)

    # 코드 블록 하단 여백
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_md_table(doc, header_row: list[str], body_rows: list[list[str]]):
    """Markdown 테이블 → Word 테이블"""
    if not header_row:
        return

    n_cols = len(header_row)
    n_rows = 1 + len(body_rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    table.autofit = True

    # 헤더 행
    for c_idx, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[c_idx]
        cell.text = cell_text.strip()
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0] if para.runs else para.add_run(cell_text.strip())
        run.font.name = '맑은 고딕'
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = COLOR_WHITE
        set_cell_bg(cell, COLOR_TABLE_HDR)
        set_cell_border(cell)

    # 데이터 행
    for r_idx, row_data in enumerate(body_rows):
        row = table.rows[r_idx + 1]
        for c_idx in range(n_cols):
            cell_text = row_data[c_idx].strip() if c_idx < len(row_data) else ""
            cell = row.cells[c_idx]
            cell.text = cell_text
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run(cell_text)
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9)
            # 짝수/홀수 행 배경
            if (r_idx % 2) == 1:
                set_cell_bg(cell, RGBColor(0xF5, 0xF8, 0xFF))
            set_cell_border(cell)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def parse_table_row(line: str) -> list[str]:
    """| col1 | col2 | → ['col1', 'col2']"""
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


def is_separator_row(cells: list[str]) -> bool:
    """| --- | :---: | 형태의 구분행 판별"""
    return all(re.match(r'^:?-+:?$', c.strip()) for c in cells if c.strip())


def convert_md_to_docx(md_path: Path, out_path: Path):
    """단일 Markdown 파일 → Word 변환"""
    text = md_path.read_text(encoding='utf-8')
    lines = text.splitlines()

    # 문서 ID와 제목 추출
    doc_id = ""
    doc_title = ""
    for line in lines[:5]:
        if line.startswith('# '):
            doc_title = line[2:].strip()
            m = re.search(r'(SF-\w+)', doc_title)
            if m:
                doc_id = m.group(1)
            break

    # 파일명에서 doc_id 추출 (fallback)
    if not doc_id:
        m = re.search(r'(SF-\w+)', md_path.stem)
        if m:
            doc_id = m.group(1)
        doc_title = md_path.stem.replace('_', ' ')

    meta = parse_meta_from_md(lines)

    doc = Document()
    set_doc_styles(doc)
    add_cover_page(doc, doc_id, doc_title, meta)
    add_header_footer(doc, doc_id, doc_title)

    # 본문 파싱
    i = 0
    in_code = False
    code_lines = []
    code_lang = ""
    table_header = []
    table_body = []
    in_table = False
    num_counter = {}

    while i < len(lines):
        line = lines[i]

        # ── 코드 블록 ─────────────────────────────────
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lang = line.strip()[3:].strip()
                code_lines = []
            else:
                add_code_block(doc, code_lines, code_lang)
                in_code = False
                code_lines = []
                code_lang = ""
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── 테이블 ────────────────────────────────────
        if '|' in line and line.strip().startswith('|'):
            cells = parse_table_row(line)
            if not in_table:
                in_table = True
                table_header = cells
                table_body = []
            elif is_separator_row(cells):
                pass  # 구분행 skip
            else:
                table_body.append(cells)
            i += 1
            # 다음 줄이 테이블이 아니면 출력
            next_line = lines[i] if i < len(lines) else ""
            if not ('|' in next_line and next_line.strip().startswith('|')):
                add_md_table(doc, table_header, table_body)
                in_table = False
                table_header = []
                table_body = []
            continue

        in_table = False

        # ── 제목 ──────────────────────────────────────
        if line.startswith('#### '):
            add_styled_heading(doc, line[5:], 4)
        elif line.startswith('### '):
            add_styled_heading(doc, line[4:], 3)
        elif line.startswith('## '):
            add_styled_heading(doc, line[3:], 2)
        elif line.startswith('# '):
            # 첫 H1은 표지에서 처리했으므로 스킵
            if i > 0:
                add_styled_heading(doc, line[2:], 1)

        # ── 수평선 ────────────────────────────────────
        elif line.strip() in ('---', '***', '___'):
            add_horizontal_rule(doc)

        # ── 불릿 ──────────────────────────────────────
        elif re.match(r'^(\s*)[-*+]\s+', line):
            indent = len(re.match(r'^(\s*)', line).group(1)) // 2
            add_bullet(doc, line, indent)

        # ── 번호 목록 ─────────────────────────────────
        elif re.match(r'^\d+\.\s+', line):
            m = re.match(r'^(\d+)\.\s+', line)
            add_numbered_item(doc, line, m.group(1) + ".")

        # ── 빈 줄 ─────────────────────────────────────
        elif line.strip() == '':
            doc.add_paragraph()

        # ── 일반 텍스트 ───────────────────────────────
        else:
            add_styled_paragraph(doc, line)

        i += 1

    doc.save(out_path)
    print(f"  ✓ {out_path.name}")


def main():
    md_files = sorted(INPUT_DIR.glob("*.md"))
    if not md_files:
        print("변환할 .md 파일이 없습니다.")
        sys.exit(1)

    print(f"\n스마트공장 공식 산출물 Word 변환 시작 ({len(md_files)}개)")
    print(f"출력 경로: {OUTPUT_DIR}\n")

    success, failed = 0, []
    for md_path in md_files:
        out_name = md_path.stem + ".docx"
        out_path = OUTPUT_DIR / out_name
        try:
            convert_md_to_docx(md_path, out_path)
            success += 1
        except Exception as e:
            print(f"  ✗ {md_path.name}: {e}")
            failed.append(md_path.name)

    print(f"\n완료: {success}개 성공" + (f", {len(failed)}개 실패: {failed}" if failed else ""))
    print(f"저장 위치: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
